from django.shortcuts import render, redirect, get_object_or_404, reverse
from django.http import HttpResponse
from django.contrib import messages
from docx import Document
import io
from django.db.models import Max

from .models import CustomerDetails, CustomerQuotation, QuotationTableInstance, QuotationTableHeader, QuotationTableValue, CustomerQuotationDetails
from .forms import CustomerForm, QuotationForm, QuotationTableForm, TableHeaderForm, TableValueForm, CustomerQuotationDetailsForm

# Step 1: Customer Information
def create_customer(request):
    if request.method == 'POST':
        customerId = request.POST.get('customer_id',0)
        quotationId = request.POST.get('quotation_id',0)
        customerIns = CustomerDetails.objects.filter(id=customerId).first() if ValueCheck(customerId) else None
        form = CustomerForm(request.POST,instance=customerIns)
        if form.is_valid():
            customer = form.save()
            request.session['customer_id'] = customer.id  # Store in session
            return redirect(reverse('add_cutomer_quotation')+ '?quotation_id='+ str(quotationId))
    else:
        customerId = request.GET.get('customer_id',0)
        customerIns = CustomerDetails.objects.filter(id=customerId).first() if ValueCheck(customerId) else None
        form = CustomerForm(instance=customerIns)
    return render(request, 'create_customer.html', {'form': form, 'customerId': customerId})

# Step 2: Quotation Name
def add_cutomer_quotation(request):
    if 'customer_id' not in request.session:
        return redirect('create_customer')
    customerId = request.session.get('customer_id') or request.GET.get('customer_id')
    if request.method == 'POST':
        quotationId = request.POST.get('quotation_id',0)
        quotationIns = CustomerQuotation.objects.get(pk=quotationId) if ValueCheck(quotationId) else None
        form = QuotationForm(request.POST,instance=quotationIns)
        if form.is_valid():
            quotationFrmIns = form.save(commit=False)
            quotationFrmIns.customer_id = customerId
            quotationFrmIns.save()
            messages.success(request, "Quotation added successfully!")
            return redirect(reverse('add_table')+ '?quotation_id='+ str(quotationFrmIns.pk)+ '&table_id=0')
    else:
        quotationId = request.GET.get('quotation_id',0)
        quotationIns = CustomerQuotation.objects.get(pk=quotationId) if ValueCheck(quotationId) else None
        form = QuotationForm(instance=quotationIns)
    quotationDetails = CustomerQuotation.objects.filter(customer_id=customerId)
    return render(request, 'add_quotation.html', {
        'form': form,
        'quotationDetails': quotationDetails,
        'quotationId': quotationId,
        'customerId': customerId
    })

def delete_quotation(request):
    if 'customer_id' not in request.session:
        return redirect('create_customer')
    quotationId = request.POST.get('quotation_id',0)
    if ValueCheck(quotationId):
        # table deletion
        isoTableInsIdList = QuotationTableInstance.objects.filter(quotation_id=quotationId).values_list('id', flat=True)
        QuotationTableValue.objects.filter(table_id__in=isoTableInsIdList).delete
        QuotationTableHeader.objects.filter(table_id__in=isoTableInsIdList).delete()
        QuotationTableInstance.objects.filter(id__in=isoTableInsIdList).delete()
        # Quotation deletion 
        quotation = get_object_or_404(CustomerQuotation, pk=quotationId)
        quotation.delete()
        messages.success(request, "Quotation deleted successfully!")
    
    # Redirect to the quotation list or another appropriate page
    return redirect(reverse('add_cutomer_quotation')+ '?quotation_id=0')

def delete_customer(request):   
    if 'customer_id' not in request.session:
        return redirect('create_customer')
    customerId = request.GET.get('customer_id',0)
    if ValueCheck(customerId):
        # Get all quotation IDs for the customer
        quotationIdList = CustomerQuotation.objects.filter(customer_id=customerId).values_list('id', flat=True)

        # Get all table instance IDs for these quotations
        isoTableInsIdList = QuotationTableInstance.objects.filter(quotation_id__in=quotationIdList).values_list('id', flat=True)

        # First delete the values that reference headers
        QuotationTableValue.objects.filter(table_id__in=isoTableInsIdList).delete()

        # Then delete the headers that reference tables
        QuotationTableHeader.objects.filter(table_id__in=isoTableInsIdList).delete()

        # Now it's safe to delete the table instances
        QuotationTableInstance.objects.filter(id__in=isoTableInsIdList).delete()

        # Delete quotation details
        CustomerQuotationDetails.objects.filter(customer_id=customerId).delete()

        # Finally delete the quotations and customer
        CustomerQuotation.objects.filter(customer_id=customerId).delete()
        CustomerDetails.objects.filter(id=customerId).delete()
        messages.success(request, "Customer deleted successfully!")
    
    # Redirect to the customer list or another appropriate page
    return redirect(reverse('all_customers'))

# Step 3: Create Table
def add_table(request):
    if 'customer_id' not in request.session:
        return redirect('create_customer')
    
    if request.method == 'POST':
        tableId = request.POST.get('table_id',0)
        quotationId = request.POST.get('quotation_id',0)
        tableIns = QuotationTableInstance.objects.get(pk=tableId) if ValueCheck(tableId) else None
        form = QuotationTableForm(request.POST,instance=tableIns)
        if form.is_valid():
            table = form.save(commit=False)
            table.quotation_id = quotationId
            table.save()
            messages.success(request, "Table Name added successfully!")
            return redirect(reverse('add_table')+ '?quotation_id='+ str(quotationId)+ '&table_id=0')
    else:
        quotationId = request.GET.get('quotation_id',0)
        tableId = request.GET.get('table_id',0)
        tableIns = QuotationTableInstance.objects.get(pk=tableId) if ValueCheck(tableId) else None
        form = QuotationTableForm(instance=tableIns)
    tableNameData = QuotationTableInstance.objects.filter(quotation_id=quotationId)
    return render(request, 'add_quotation_table.html', {
        'forms': form,
        'tableNameData': tableNameData,
        'quotationId': quotationId,
        'tableId': tableId,
    })

# Step 4: Create Table Headers
def add_table_headers(request):
    if 'customer_id' not in request.session:
        return redirect('create_customer')
    
    if request.method == 'POST':
        tableId = request.POST.get('table_id',0)
        fieldId = request.POST.get('field_id',0)
        quotationId = request.POST.get('quotation_id',0)
        isoTableHeaderIns = QuotationTableHeader.objects.get(pk=fieldId) if ValueCheck(fieldId) else None
        form = TableHeaderForm(request.POST, instance=isoTableHeaderIns)
        if form.is_valid():
            tableHeaderIns = form.save(commit=False)
            tableHeaderIns.table_id = tableId
            tableHeaderIns.save()
            messages.success(request, "Table Headers added successfully!")
            return redirect(reverse('add_table_headers')+ '?quotation_id='+ str(quotationId)+ '&table_id='+str(tableId)+'')
    else:
        tableId = request.GET.get('table_id',0)
        fieldId = request.GET.get('field_id',0)
        quotationId = request.GET.get('quotation_id',0)
        isoTableHeaderIns = QuotationTableHeader.objects.get(pk=fieldId) if ValueCheck(fieldId) else None
        form = TableHeaderForm(instance=isoTableHeaderIns)
    isoAmountFieldId = QuotationTableHeader.objects.filter(field_type__field_type='Amount').exists()
    tableHeadersData = QuotationTableHeader.objects.filter(table_id=tableId)
    
    return render(request, 'add_table_headers.html', {
        'forms': form,
        'tableHeadersData': tableHeadersData,
        'quotationId': quotationId,
        'tableId': tableId,
        'fieldId': fieldId,
        "isoAmountFieldId": isoAmountFieldId
    })

# Step 5: Create Table Values
def add_table_value(request):
    if 'customer_id' not in request.session:
        return redirect('create_customer')
    if request.method == 'POST':
        tableRowId = request.POST.get('row_id',0)
        tableId = request.POST.get('table_id',0)
        fieldId = request.POST.get('field_id',0)
        quotationId = request.POST.get('quotation_id',0)
        isoFieldValueList  = request.POST.getlist('DynamicValues')
        isoFieldIdList  = request.POST.getlist('DynamicHiddenIds')
        tableHeaders = QuotationTableHeader.objects.filter(table_id=tableId)
        tableFieldIdList = tableHeaders.values_list('id', flat=True)   
        isoTableFieldValuesData = QuotationTableValue.objects.filter(field_id__in=tableFieldIdList)
        MaxOrder = isoTableFieldValuesData.aggregate(Max('row_id')).get('row_id__max')  or 0   
        isoTableFieldValues = isoTableFieldValuesData.filter(row_id=tableRowId) 
        tableTotalvalue = QuotationTableInstance.objects.get(id=tableId).table_total or 0.0 if ValueCheck(tableId) else 0.0
        isoAmountFieldId = tableHeaders.get(field_type__field_type='Amount').id
        for isoFieldId, isoFieldValue in zip(isoFieldIdList, isoFieldValueList):
            if isoFieldValue:
                if isoTableFieldValues.exists():
                    isoOldData = QuotationTableValue.objects.filter(field_id=isoFieldId, row_id=tableRowId).exists()
                    if isoOldData:
                        isoOldvalue = QuotationTableValue.objects.get(field_id=isoFieldId, row_id=tableRowId).field_value
                        if isoOldvalue != isoFieldValue:
                            QuotationTableValue.objects.filter(field_id=isoFieldId, row_id=tableRowId).update(field_value=isoFieldValue)
                            if int(isoFieldId) == isoAmountFieldId:
                                tableTotalvalue += float(int(isoFieldValue))
                                QuotationTableInstance.objects.filter(id=tableId).update(table_total=tableTotalvalue)
                    else:
                        QuotationTableValue.objects.create(field_id=isoFieldId, field_value=isoFieldValue, row_id=tableRowId,table_id=tableId)
                        if int(isoFieldId) == isoAmountFieldId:
                            tableTotalvalue += float(int(isoFieldValue))
                            QuotationTableInstance.objects.filter(id=tableId).update(table_total=tableTotalvalue)
                else:
                    QuotationTableValue.objects.create(field_id=isoFieldId, field_value=isoFieldValue, row_id=MaxOrder+1,table_id=tableId)
                    if int(isoFieldId) == isoAmountFieldId:
                        tableTotalvalue += float(int(isoFieldValue))
                        QuotationTableInstance.objects.filter(id=tableId).update(table_total=tableTotalvalue)
        messages.success(request, "Table Value added successfully!")
        return redirect(reverse('add_table_value')+ '?quotation_id='+ str(quotationId)+ '&table_id='+str(tableId)+'')
    else:
        tableId = request.GET.get('table_id',0)
        fieldId = request.GET.get('field_id',0)
        tableRowId = request.GET.get('row_id',0)
        quotationId = request.GET.get('quotation_id',0)
        form = TableValueForm(tableId=tableId,tableRowId=tableRowId,datalists=0)
        tableHeaders = QuotationTableHeader.objects.filter(table_id=tableId)
        tableValues = QuotationTableValue.objects.filter(table_id=tableId,field_id__in=tableHeaders.values_list('id', flat=True))
    
    return render(request, 'add_table_data.html', {
        'forms': form,
        'tableHeaders': tableHeaders,
        'tableValues':tableValues,
        'quotationId': quotationId,
        'tableId': tableId,
        'fieldId': fieldId,
        'tableRowId': tableRowId
        
    })

# Step 4: Complaint Details
def add_complaint_details(request):
    if 'customer_id' not in request.session:
        return redirect('create_customer')
    isoContext={}
    customerId = request.session.get('customer_id') or request.GET.get('customer_id')
    if request.method == 'POST':
        quotationId = request.POST.get('quotation_id', 0)
        # customerDeatilsId = request.POST.get('customer_details_id')
        customerQuotationIns = CustomerQuotationDetails.objects.get(customer_id=customerId, quotation_id=quotationId)
        form = CustomerQuotationDetailsForm(request.POST,instance=customerQuotationIns)
        if form.is_valid():
            complaint = form.save(commit=False)
            complaint.quotation_id = quotationId
            complaint.customer_id = customerId
            complaint.save()
            # del request.session['customer_id']  # Cleanup
            return redirect('document_preview', quotation_id=quotationId)
    else:
        quotationId = request.GET.get('quotation_id', 0)
        customerQuotationIns = CustomerQuotationDetails.objects.get(customer_id=customerId, quotation_id=quotationId) if CustomerQuotationDetails.objects.filter(customer_id=customerId, quotation_id=quotationId).exists() else None 
        form = CustomerQuotationDetailsForm(instance=customerQuotationIns)
        tableId = request.GET.get('table_id',0)
        customerIns = CustomerDetails.objects.filter(id=customerId)
        quotationIns = CustomerQuotation.objects.get(pk=quotationId)
        tableHeaders = QuotationTableHeader.objects.filter(table_id=tableId)
        tableValues = QuotationTableValue.objects.filter(table_id=tableId,field_id__in=tableHeaders.values_list('id', flat=True))
        tableheadervalues = view_quotation_tables(quotationId)
        isoContext |= tableheadervalues
        isoContext['form'] = form
        isoContext['customer'] = customerIns
        isoContext['customerIns'] = customerIns
        isoContext['quotationIns'] = quotationIns
        isoContext['tableHeaders'] = tableHeaders
        isoContext['tableValues'] = tableValues
        isoContext['quotationId'] = quotationId
        isoContext['tableId'] = tableId
    return render(request, 'add_complaint_details.html', isoContext)

from django.db.models import Sum
import re
def document_preview(request):
    # Get the complaint and related data
    customerId = request.session.get('customer_id') or request.GET.get('customer_id')
    quotationId = request.GET.get('quotation_id', 0)
    isoContext={}
    quotationDetails = CustomerQuotation.objects.filter(customer_id=customerId)
    if quotationId:
        quotationDetails = quotationDetails.filter(id=quotationId)
        customerIns = CustomerDetails.objects.get(id=customerId)
        quotationIns = CustomerQuotation.objects.get(pk=quotationId)
        tableIdList = QuotationTableInstance.objects.filter(quotation_id=quotationId).values_list('id', flat=True)
        tableHeaders = QuotationTableHeader.objects.filter(table_id__in=tableIdList)
        tableValues = QuotationTableValue.objects.filter(table_id__in=tableIdList,field_id__in=tableHeaders.values_list('id', flat=True))
        tableheadervalues = view_quotation_tables(quotationId)
        customerQuotationIns = CustomerQuotationDetails.objects.get(customer_id=customerId, quotation_id=quotationId)
        grandTotal = QuotationTableInstance.objects.filter(quotation_id=quotationId).aggregate(total=Sum('table_total'))['total'] or 0.00
        isoContext['customer'] = customerIns
        isoContext['customerDeatils'] = customerQuotationIns
        isoContext['customerIns'] = customerIns
        isoContext['quotationIns'] = quotationIns
        isoContext['tableHeaders'] = tableHeaders
        isoContext['tableValues'] = tableValues
        isoContext['quotationId'] = quotationId
        isoContext['grandTotal'] = grandTotal
        isoContext |= tableheadervalues
        return render(request, 'document_preview.html', isoContext)
    return render(request, 'all_quotation.html', {'quotationDetails': quotationDetails})

from docx import Document
from docx.shared import Pt, Inches,RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.http import HttpResponse
from io import BytesIO
import html

def download_quotation_doc(request, quotation_id):
    # Get your data
    customer_id = request.session.get('customer_id') or request.GET.get('customer_id')
    customer_details = CustomerQuotationDetails.objects.get(
        customer_id=customer_id, 
        quotation_id=quotation_id
    )
    quotationDetails = CustomerQuotation.objects.get(id=quotation_id)
    customer = CustomerDetails.objects.get(id=customer_id)
    table_data = view_quotation_tables(quotation_id)
    grand_total = QuotationTableInstance.objects.filter(
        quotation_id=quotation_id
    ).aggregate(total=Sum('table_total'))['total'] or 0.00

    regexvalue = re.compile('<.*?>') 


    # Create document
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    font.color.rgb = RGBColor(0, 0, 0)  # Force black color for all text

    # Header section - remove <br> tags by using proper paragraph structure
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_para.add_run(f"{customer_details.quotation_date}\n")
    
    to_para = doc.add_paragraph()
    to_run = to_para.add_run("To:")
    to_run.bold = True
    
    customer_para = doc.add_paragraph(customer.customer_name)
    address_para = doc.add_paragraph(customer_details.quotation_address)

    # Subject
    subject_para = doc.add_paragraph()
    subject_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subjectDetails = html.unescape(re.sub(regexvalue, '', customer_details.quotation_subject))
    subject_run = subject_para.add_run(subjectDetails)
    subject_run.bold = True
    subject_run.underline = True

    # Description
    # doc.add_paragraph("Description:", style='Heading 2')
    DescriptionDetails = html.unescape(re.sub(regexvalue, '', customer_details.quotation_description))
    doc.add_paragraph(DescriptionDetails)

    # Tables
    for table in table_data['tables_data']:
        doc.add_paragraph(table['table'].table_name, style='Heading 3')
        
        # Create table with headers and set column widths
        word_table = doc.add_table(rows=1, cols=len(table['headers']) + 1)
        word_table.style = 'Table Grid'
        
        # Set column widths - S.no column narrower
        if len(word_table.columns) > 0:
            word_table.columns[0].width = Inches(0.5)  # Narrow width for S.no
        
        hdr_cells = word_table.rows[0].cells
        hdr_cells[0].text = 'S.no'
        
        # Set headers
        for i, header in enumerate(table['headers'], 1):
            hdr_cells[i].text = header.field_name
            if header.field_type_id == 2:
                hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add rows
        for row_idx, row in enumerate(table['rows'], 1):
            row_cells = word_table.add_row().cells
            row_cells[0].text = str(row_idx)
            
            for i, header in enumerate(table['headers'], 1):
                value = row.get(header.id, "")
                cell = row_cells[i]
                
                if header.field_type_id == 2:
                    cell.text = f"{float(value or 0):.2f}"
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    cell.text = str(value)
        
        # Add totals
        total_row = word_table.add_row()
        total_cells = total_row.cells
        total_cells[0].merge(total_cells[len(table['headers']) - 1])
        total_paragraph = total_cells[0].paragraphs[0]
        total_run = total_paragraph.add_run("Total")
        total_run.bold = True
        total_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        total_cells[-1].text = f"{float(table['table'].table_total or 0):.2f}"
        total_cells[-1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Grand total for last table
        if table == table_data['tables_data'][-1]:
            grand_row = word_table.add_row()
            grand_cells = grand_row.cells
            grand_cells[0].merge(grand_cells[len(table['headers']) - 1])
            grand_total_paragraph = grand_cells[0].paragraphs[0]
            grand_run = grand_total_paragraph.add_run("Grand Total")
            grand_run.bold = True
            grand_total_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            grand_cells[-1].text = f"{float(grand_total or 0):.2f}"
            grand_cells[-1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Footer sections - remove <br> tags
    referrenceDetails = html.unescape(re.sub(regexvalue, '', customer_details.quotation_referrence))
    doc.add_paragraph(referrenceDetails)
    
    doc.add_paragraph("Note:", style='Heading 2')
    noteDetails = html.unescape(re.sub(regexvalue, '', customer_details.quotation_note))
    doc.add_paragraph(noteDetails)
    
    doc.add_paragraph("Bank Details:", style='Heading 2')
    bankDetails = html.unescape(re.sub(regexvalue, '', customer_details.quotation_bank_details))
    doc.add_paragraph(bankDetails)
    
    # Signature
    sign_para = doc.add_paragraph()
    sign_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    signatureDetails = html.unescape(re.sub(regexvalue, '', customer_details.quotation_signature))
    sign_run = sign_para.add_run(signatureDetails)
    sign_run.bold = True

    # Generate response
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    response = HttpResponse(
        file_stream.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename={quotationDetails.quotation_name}.docx'
    return response

def generate_doc(request, complaint_id):
    # Get the complaint and related data
    complaint = get_object_or_404(CustomerDetails, pk=complaint_id)
    customer = complaint.customer
    headers = QuotationTableHeader.objects.filter(complaint=complaint)
    
    # Create a new Word document
    doc = Document()
    
    # Add document header
    doc.add_heading('Complaint Document', level=1)
    
    # Add customer information
    doc.add_paragraph(f"Customer: {customer.name}")
    doc.add_paragraph(f"Address: {customer.address}")
    
    # Add complaint details
    doc.add_heading('Complaint Details', level=2)
    doc.add_paragraph(f"Date: {complaint.date}")
    doc.add_paragraph(f"Subject: {complaint.subject}")
    doc.add_paragraph(f"Description: {complaint.description}")
    
    # Add tables
    table_ids = set(headers.values_list('table_id', flat=True))
    for table_id in table_ids:
        table_headers = headers.filter(table_id=table_id)
        table_data = QuotationTableValue.objects.filter(header__in=table_headers)
        
        # Create table in document
        table = doc.add_table(rows=1, cols=len(table_headers))
        table.style = 'Table Grid'
        
        # Add headers
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(table_headers):
            hdr_cells[i].text = header.name
        
        # Add data rows
        # Group data by row - this is simplified
        for i in range(len(table_data) // len(table_headers)):
            row_cells = table.add_row().cells
            for j, header in enumerate(table_headers):
                try:
                    data = table_data.filter(header=header)[i]
                    row_cells[j].text = data.value
                except IndexError:
                    row_cells[j].text = ""
    
    # Add additional details
    doc.add_heading('Additional Information', level=2)
    doc.add_paragraph(f"Note: {complaint.note}")
    doc.add_paragraph(f"Bank Details: {complaint.bank_details}")
    doc.add_paragraph(f"Signature: {complaint.signature}")
    
    # Save document to a BytesIO stream
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Create HTTP response
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename=complaint_{complaint_id}.docx'
    return response


def all_customers(request):
    customer = CustomerDetails.objects.all()
    return render(request, 'all_customer.html', {'customer': customer})

def ValueCheck(isoVal):
    isoVal = isoVal.strip() if isinstance(isoVal, str) else isoVal
    return bool(isoVal) and isoVal not in ('0', 'None', [''],'[]')

def view_quotation_tables(quotation_id):
    tables = QuotationTableInstance.objects.filter(quotation_id=quotation_id)
    
    tables_data = []
    for table in tables:
        headers = QuotationTableHeader.objects.filter(table=table).order_by('id')
        row_ids = QuotationTableValue.objects.filter(
            field__table=table
        ).values_list('row_id', flat=True).distinct()
        
        # Create a list of dictionaries with header.id as keys and direct values
        rows = []
        for row_id in row_ids:
            row_values = QuotationTableValue.objects.filter(
                field__table=table,
                row_id=row_id
            ).select_related('field')
            
            # Create a simple dictionary with header.id as keys and direct values
            row_data = {'row_id': row_id}
            for value in row_values:
                row_data[value.field_id] = value.field_value  # Store direct value
            rows.append(row_data)
        
        tables_data.append({
            'table': table,
            'headers': headers,
            'rows': rows
        })
    response_data = {
        'tables_data': tables_data,
        'quotation_id': quotation_id
    }
    return response_data